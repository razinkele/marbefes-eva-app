"""Generate improved PA report using EUNIS L3 data in BBT8 format.

Produces both HTML (interactive maps) and Word (static) reports
using the pre-extracted EUNIS overlay and corrected EVA data.
"""
import logging
import os
import sys
import io
from datetime import datetime

import geopandas as gpd
import numpy as np
import pandas as pd
import folium
import folium.plugins
import branca.colormap as cm
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# Ensure project root on path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import eunis_data
from pa_export import generate_bbt8_workbook

TUTORIAL_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "tutorial"))
CORRECTED_DIR = os.path.normpath(
    r"C:\Users\DELL\OneDrive - ku.lt\HORIZON_EUROPE\MARBEFES\EVA_FINAL_corrected"
)

EUNIS_COLORS = {
    "A5.26 or A5.35 or A5.36": "#8b7355",
    "A5.23": "#ffe4b5",
    "A5.25": "#f4a460",
    "A5.14": "#d2b48c",
    "A4.4": "#2e8b57",
    "A3.4": "#006400",
    "A5.24 or A5.33 or A5.34": "#cd853f",
    "A5.27 or A5.37": "#a0522d",
    "A5.13": "#deb887",
}
DEFAULT_COLOR = "#999999"


def load_all():
    logger.info("Loading data...")
    eunis_path = os.path.join(TUTORIAL_DIR, "eunis_l3_lithuanian.gpkg")
    overlay = eunis_data.load_eunis_overlay(eunis_path)
    logger.info("  EUNIS overlay: %d hexes, %d with data", len(overlay), overlay["dominant_EUNIS"].notna().sum())

    eva_path = os.path.join(CORRECTED_DIR, "ALL4EVA_2025_fixed_geometries.gpkg")
    eva = gpd.read_file(eva_path)
    logger.info("  EVA: %d features", len(eva))

    return overlay, eva


def compute_all(overlay, eva):
    logger.info("Computing accounts...")
    extent = eunis_data.compute_eunis_extent(overlay, unit="Ha")
    condition = eunis_data.compute_eunis_condition(overlay, eva)
    supply = eunis_data.compute_eunis_supply(overlay, eva)
    accounts = eunis_data.build_accounts_summary(extent, condition)
    missing = eunis_data.build_missing_values(overlay, eva, total_bbt_area_m2=0)
    return extent, condition, supply, accounts, missing


def make_eunis_map_folium(overlay):
    """Interactive folium map of EUNIS types."""
    gdf = overlay.to_crs(epsg=4326)
    bounds = gdf.total_bounds
    center = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
    m = folium.Map(location=center, zoom_start=9, tiles="cartodbpositron")

    def style_fn(feature):
        code = feature["properties"].get("dominant_EUNIS", "")
        return {
            "fillColor": EUNIS_COLORS.get(code, DEFAULT_COLOR),
            "color": "#333", "weight": 0.5,
            "fillOpacity": 0.7 if code else 0.2,
        }

    gdf_plot = gdf[["Subzone_ID", "dominant_EUNIS", "dominant_EUNIS_name",
                     "habitat_count", "dominant_pct", "geometry"]].copy()
    gdf_plot["dominant_pct"] = gdf_plot["dominant_pct"].round(1)

    folium.GeoJson(
        gdf_plot.to_json(),
        style_function=style_fn,
        tooltip=folium.GeoJsonTooltip(
            fields=["Subzone_ID", "dominant_EUNIS", "dominant_EUNIS_name", "habitat_count", "dominant_pct"],
            aliases=["Subzone:", "EUNIS Code:", "Habitat:", "Types:", "Dominant %:"],
        ),
    ).add_to(m)

    legend = '<div style="position:fixed;bottom:30px;left:30px;z-index:1000;background:white;padding:12px;border-radius:8px;box-shadow:0 2px 6px rgba(0,0,0,0.3);font-size:11px;max-height:300px;overflow-y:auto;">'
    legend += '<strong>EUNIS Level 3</strong><br>'
    for code, color in EUNIS_COLORS.items():
        short = code[:20] + "..." if len(code) > 20 else code
        legend += f'<span style="background:{color};width:12px;height:12px;display:inline-block;margin-right:4px;border-radius:2px;"></span>{short}<br>'
    legend += '</div>'
    m.get_root().html.add_child(folium.Element(legend))
    folium.plugins.Fullscreen().add_to(m)
    m.fit_bounds([[bounds[1], bounds[0]], [bounds[3], bounds[2]]])
    return m._repr_html_()


def make_condition_map_folium(overlay, eva, variable, title):
    """Interactive condition map."""
    gdf = overlay.copy()
    eva_sub = eva[["Subzone_ID", variable]].copy() if variable in eva.columns else None
    if eva_sub is not None and "geometry" in eva_sub.columns:
        eva_sub = eva_sub.drop(columns="geometry")
    if eva_sub is not None:
        gdf = gdf.merge(eva_sub, on="Subzone_ID", how="left")
    else:
        gdf[variable] = np.nan

    gdf = gdf.to_crs(epsg=4326)
    bounds = gdf.total_bounds
    center = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
    m = folium.Map(location=center, zoom_start=9, tiles="cartodbpositron")

    vals = gdf[variable].dropna()
    vmin, vmax = (float(vals.min()), float(vals.max())) if len(vals) > 0 else (0, 5)
    if vmax == vmin:
        vmax = vmin + 1
    colormap = cm.linear.RdYlGn_11.scale(vmin, vmax)
    colormap.caption = title

    gdf[variable] = gdf[variable].round(3)

    def style_fn(feature):
        val = feature["properties"].get(variable)
        if val is None:
            return {"fillColor": "#ccc", "color": "#333", "weight": 0.5, "fillOpacity": 0.3}
        return {"fillColor": colormap(val), "color": "#333", "weight": 0.5, "fillOpacity": 0.7}

    folium.GeoJson(
        gdf[["Subzone_ID", "dominant_EUNIS", variable, "geometry"]].to_json(),
        style_function=style_fn,
        tooltip=folium.GeoJsonTooltip(
            fields=["Subzone_ID", "dominant_EUNIS", variable],
            aliases=["Subzone:", "EUNIS:", f"{title}:"],
        ),
    ).add_to(m)
    colormap.add_to(m)
    folium.plugins.Fullscreen().add_to(m)
    m.fit_bounds([[bounds[1], bounds[0]], [bounds[3], bounds[2]]])
    return m._repr_html_()


def make_eunis_map_matplotlib(overlay):
    """Static matplotlib map."""
    gdf = overlay.to_crs(epsg=4326)
    fig, ax = plt.subplots(figsize=(10, 8))
    for code, color in EUNIS_COLORS.items():
        sub = gdf[gdf["dominant_EUNIS"] == code]
        if not sub.empty:
            sub.plot(ax=ax, color=color, edgecolor="#333", linewidth=0.3, label=code)
    no_data = gdf[gdf["dominant_EUNIS"].isna()]
    if not no_data.empty:
        no_data.plot(ax=ax, color="#eee", edgecolor="#999", linewidth=0.3, label="No data")
    ax.set_title("EUNIS Level 3 Habitat Classification — Lithuanian BBT5",
                 fontsize=14, fontweight="bold", color="#006994")
    ax.set_xlabel("Longitude"); ax.set_ylabel("Latitude")
    patches = [mpatches.Patch(color=c, label=n) for n, c in EUNIS_COLORS.items()]
    ax.legend(handles=patches, loc="lower left", fontsize=7, framealpha=0.9)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def make_condition_map_matplotlib(overlay, eva, variable, title):
    """Static condition map."""
    gdf = overlay.copy()
    eva_sub = eva[["Subzone_ID", variable]].copy() if variable in eva.columns else None
    if eva_sub is not None and "geometry" in eva_sub.columns:
        eva_sub = eva_sub.drop(columns="geometry")
    if eva_sub is not None:
        gdf = gdf.merge(eva_sub, on="Subzone_ID", how="left")
    else:
        gdf[variable] = np.nan
    gdf = gdf.to_crs(epsg=4326)
    fig, ax = plt.subplots(figsize=(10, 8))
    gdf.plot(column=variable, ax=ax, cmap="RdYlGn", vmin=0, vmax=5,
             edgecolor="#333", linewidth=0.3, legend=True, missing_kwds={"color": "#eee"},
             legend_kwds={"label": f"{title} (0-5)", "shrink": 0.7})
    ax.set_title(f"{title} — Lithuanian BBT5", fontsize=14, fontweight="bold", color="#006994")
    ax.set_xlabel("Longitude"); ax.set_ylabel("Latitude")
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def make_extent_chart_matplotlib(extent):
    fig, ax = plt.subplots(figsize=(10, 5))
    codes = extent["EUNIS_code"].tolist()
    areas = extent["total_area"].tolist()
    colors = [EUNIS_COLORS.get(c, DEFAULT_COLOR) for c in codes]
    bars = ax.barh(codes, areas, color=colors, edgecolor="#333", linewidth=0.5)
    ax.set_xlabel("Area (Ha)", fontsize=12)
    ax.set_title("Ecosystem Extent by EUNIS L3 Habitat Type", fontsize=14, fontweight="bold", color="#006994")
    for bar, val in zip(bars, areas):
        ax.text(bar.get_width() + max(areas)*0.01, bar.get_y() + bar.get_height()/2,
                f"{val:,.0f} Ha", va="center", fontsize=9)
    ax.invert_yaxis()
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def make_condition_heatmap_matplotlib(condition):
    indicators = ["Habitat_EV", "Habitat_confidence"]
    avail = [c for c in indicators if c in condition.columns]
    if not avail:
        return None
    data = condition[avail].values
    fig, ax = plt.subplots(figsize=(8, 5))
    im = ax.imshow(data, cmap="RdYlGn", vmin=0, vmax=5, aspect="auto")
    ax.set_xticks(range(len(avail)))
    ax.set_xticklabels(["Habitat EV", "Confidence"][:len(avail)], fontsize=11)
    ax.set_yticks(range(len(condition)))
    ax.set_yticklabels(condition["EUNIS_code"], fontsize=9)
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            v = data[i, j]
            if not np.isnan(v):
                ax.text(j, i, f"{v:.2f}", ha="center", va="center", fontsize=10,
                        color="white" if v < 1.5 or v > 3.5 else "black")
    plt.colorbar(im, ax=ax, label="Score (0-5)", shrink=0.8)
    ax.set_title("Condition per EUNIS Class", fontsize=14, fontweight="bold", color="#006994")
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def classify_eva(val):
    if pd.isna(val): return "No Data"
    if val <= 1: return "Very Low"
    if val <= 2: return "Low"
    if val <= 3: return "Medium"
    if val <= 4: return "High"
    return "Very High"


def df_to_html(df):
    html = '<table style="width:100%;border-collapse:collapse;font-size:13px;margin:1em 0;">'
    html += '<thead><tr style="background:linear-gradient(135deg,#006994,#00b8d4);color:white;">'
    for col in df.columns:
        html += f'<th style="padding:8px 10px;text-align:left;border-bottom:2px solid #ddd;">{col}</th>'
    html += '</tr></thead><tbody>'
    for i, (_, row) in enumerate(df.iterrows()):
        bg = "#f8f9fa" if i % 2 == 0 else "white"
        html += f'<tr style="background:{bg};">'
        for col in df.columns:
            html += f'<td style="padding:6px 10px;border-bottom:1px solid #eee;">{row[col]}</td>'
        html += '</tr>'
    html += '</tbody></table>'
    return html


def generate_html_report(overlay, eva, extent, condition, supply, accounts, missing):
    logger.info("Generating HTML report...")
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Tables
    ext_disp = extent[["EUNIS_code", "EUNIS_name", "n_subzones", "total_area", "pct_of_total"]].copy()
    ext_disp.columns = ["EUNIS Code", "Habitat", "Subzones", "Area (Ha)", "% Total"]

    cond_disp = condition.copy()
    cond_disp["EV Class"] = cond_disp["Habitat_EV"].apply(classify_eva)
    cond_disp = cond_disp.rename(columns={
        "EUNIS_code": "EUNIS Code", "EUNIS_name": "Habitat",
        "Habitat_EV": "Habitat EV", "Habitat_confidence": "Confidence",
        "n_subzones": "Subzones",
    })

    supply_disp = supply.rename(columns={
        "EUNIS_code": "EUNIS Code", "EUNIS_name": "Habitat",
        "Fisheries_proxy": "Fisheries", "FoodWeb_proxy": "Food Web",
        "PrimaryProd_proxy": "Primary Prod.",
    })

    acct_disp = accounts.copy()
    acct_disp["area_m2"] = acct_disp["area_m2"].apply(lambda x: f"{x:,.0f}")
    acct_disp["Habitat_EV"] = acct_disp["Habitat_EV"].round(2)
    acct_disp["Habitat_confidence"] = acct_disp["Habitat_confidence"].round(2)
    acct_disp.columns = ["EUNIS Code", "Habitat", "Area (m2)", "Habitat EV", "Confidence"]

    # Maps
    habitat_map = make_eunis_map_folium(overlay)
    condition_maps = {}
    for var, name in [("TotalEV_MAX", "Total EV"), ("MaxBenthos", "Benthos Condition"),
                      ("ZooScore", "Zooplankton"), ("PhytoScore", "Phytoplankton")]:
        if var in eva.columns:
            condition_maps[name] = make_condition_map_folium(overlay, eva, var, name)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>BBT8 Physical Accounts — Lithuanian BBT5 (EUNIS L3)</title>
<style>
  body {{ font-family:'Segoe UI',sans-serif; margin:0; padding:0; background:#f5f7fa; color:#333; }}
  .container {{ max-width:1200px; margin:0 auto; padding:20px; }}
  .header {{ background:linear-gradient(135deg,#006994,#00b8d4); color:white; padding:40px; border-radius:12px; margin-bottom:30px; }}
  .header h1 {{ margin:0; font-size:2.2em; }}
  .header p {{ margin:8px 0 0; opacity:0.9; }}
  .section {{ background:white; border-radius:12px; padding:30px; margin-bottom:25px; box-shadow:0 2px 8px rgba(0,0,0,0.08); }}
  .section h2 {{ color:#006994; border-bottom:3px solid #00b8d4; padding-bottom:10px; margin-top:0; }}
  .section h3 {{ color:#006994; }}
  .map-container {{ border-radius:8px; overflow:hidden; border:1px solid #ddd; margin:15px 0; }}
  .info-box {{ background:linear-gradient(135deg,#e3f2fd,#bbdefb); border-left:4px solid #006994; padding:15px; border-radius:8px; margin:15px 0; }}
  .warn-box {{ background:linear-gradient(135deg,#fff3e0,#ffe0b2); border-left:4px solid #ff9800; padding:15px; border-radius:8px; margin:15px 0; }}
  .footer {{ text-align:center; color:#6c757d; padding:20px; font-size:0.9em; }}
</style>
</head>
<body>
<div class="container">

<div class="header">
  <h1>SEEA EA Physical Accounts (BBT8 Format)</h1>
  <p>Lithuanian BBT5 — Curonian Lagoon and Baltic Sea Coast</p>
  <p>EUNIS Level 3 Classification from EMODnet EUSeaMap 2023</p>
  <p style="font-size:0.85em;opacity:0.7;">MARBEFES WP4 | Generated: {now}</p>
</div>

<div class="section">
  <h2>1. Introduction</h2>
  <p>This report presents SEEA EA Physical Natural Capital Accounts for the Lithuanian BBT5
  using <strong>EUNIS Level 3</strong> habitat classification from the EMODnet EUSeaMap 2023.
  This replaces the earlier MSFD broad-type classification with the standardized
  EUNIS system used across all MARBEFES BBTs.</p>
  <div class="info-box">
    <strong>Spatial coverage:</strong> 425 hexagonal subzones (3 km grid), 391 with EUNIS data (92%)<br>
    <strong>Habitat types:</strong> 9 EUNIS Level 3 classes identified<br>
    <strong>EVA data:</strong> Sentinel-corrected, MAX-aggregated (September 2025)
  </div>
</div>

<div class="section">
  <h2>2. Extent Account</h2>
  <h3>2.1 Area per EUNIS Class</h3>
  {df_to_html(ext_disp)}

  <h3>2.2 Habitat Distribution Map</h3>
  <div class="map-container" style="height:600px;">
    {habitat_map}
  </div>
</div>

<div class="section">
  <h2>3. Condition Account</h2>
  <div class="warn-box">
    <strong>Confidence:</strong> All scores are Low confidence (0.3) — each EC answered only 1-4 of 7-8 possible AQs.
  </div>
  <h3>3.1 Condition per EUNIS Class</h3>
  {df_to_html(cond_disp)}
"""

    for name, map_html in condition_maps.items():
        html += f"""
  <h3>{name}</h3>
  <div class="map-container" style="height:500px;">{map_html}</div>
"""

    html += f"""
</div>

<div class="section">
  <h2>4. Supply Table (Proxy Scores)</h2>
  <p>Ecosystem service proxies based on EVA scores (0-5 scale). Physical units required for full SEEA EA compliance.</p>
  {df_to_html(supply_disp)}
</div>

<div class="section">
  <h2>5. BBT8 Accounts Summary</h2>
  <p>Standard accounts table following the Irish Sea BBT8 format.</p>
  {df_to_html(acct_disp)}
</div>

<div class="section">
  <h2>6. Data Quality</h2>
  <p>Missing values: {len(missing)} subzones with issues.</p>
  {"<p>No missing values detected.</p>" if missing.empty else df_to_html(missing.head(20))}
</div>

<div class="section">
  <h2>References</h2>
  <ul>
    <li>Franco A. &amp; Amorim E. (2025) EVA Guidance. MARBEFES WP4.1.</li>
    <li>EMODnet (2023) EUSeaMap 2023 — European marine habitat map.</li>
    <li>UN (2021) SEEA EA — System of Environmental-Economic Accounting.</li>
    <li>Razinkovas-Baziukas A. et al. (2025) Lithuanian BBT5 EVA report.</li>
  </ul>
</div>

<div class="footer">
  MARBEFES — EU Horizon Europe | EVA v3.4.0 | {now}
</div>
</div>
</body>
</html>"""

    path = os.path.join(TUTORIAL_DIR, "PhysicalAccounts_BBT8_LithuanianBBT5.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    logger.info("HTML report: %s", path)
    return path


def generate_docx_report(overlay, eva, extent, condition, supply, accounts, missing):
    """Generate Word document version."""
    logger.info("Generating Word report...")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    now = datetime.now().strftime("%Y-%m-%d")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def set_shading(cell, color):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)

    def add_table(doc, df):
        t = doc.add_table(rows=1+len(df), cols=len(df.columns))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, col in enumerate(df.columns):
            c = t.rows[0].cells[j]
            c.text = str(col)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)
            set_shading(c, "006994")
        for i, (_, row) in enumerate(df.iterrows()):
            for j, col in enumerate(df.columns):
                c = t.rows[i+1].cells[j]
                v = row[col]
                c.text = str(v) if pd.notna(v) else ""
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: r.font.size = Pt(9)
                if i % 2 == 0: set_shading(c, "F0F7FA")

    # Title
    doc.add_paragraph()
    h = doc.add_heading("SEEA EA Physical Accounts (BBT8 Format)", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs: r.font.color.rgb = RGBColor(0,105,148)
    s = doc.add_paragraph("Lithuanian BBT5 — EUNIS Level 3 (EUSeaMap 2023)")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(16); s.runs[0].font.color.rgb = RGBColor(0,184,212)
    doc.add_paragraph(f"MARBEFES WP4 | {now}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for item in ["1. Introduction", "2. Extent Account", "3. Condition Account",
                  "4. Supply Table", "5. BBT8 Accounts Summary", "6. Data Quality", "References"]:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This report presents SEEA EA Physical Accounts for Lithuanian BBT5 using "
        "EUNIS Level 3 habitat classification from EMODnet EUSeaMap 2023. "
        "9 habitat types identified across 391 of 425 hexagonal subzones (92% coverage)."
    )
    doc.add_page_break()

    # 2. Extent
    doc.add_heading("2. Extent Account", level=1)
    ext = extent[["EUNIS_code","EUNIS_name","n_subzones","total_area","pct_of_total"]].copy()
    ext.columns = ["EUNIS Code","Habitat","Subzones","Area (Ha)","% Total"]
    add_table(doc, ext)

    logger.info("  Generating extent chart...")
    doc.add_paragraph()
    doc.add_picture(make_extent_chart_matplotlib(extent), width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    logger.info("  Generating habitat map...")
    doc.add_paragraph()
    doc.add_picture(make_eunis_map_matplotlib(overlay), width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 1. EUNIS Level 3 habitat distribution.").italic = True
    doc.add_page_break()

    # 3. Condition
    doc.add_heading("3. Condition Account", level=1)
    cond = condition.copy()
    cond["EV Class"] = cond["Habitat_EV"].apply(classify_eva)
    cond = cond.rename(columns={"EUNIS_code":"EUNIS Code","EUNIS_name":"Habitat",
                                "Habitat_EV":"Habitat EV","Habitat_confidence":"Confidence",
                                "n_subzones":"Subzones"})
    add_table(doc, cond)

    logger.info("  Generating condition heatmap...")
    hm = make_condition_heatmap_matplotlib(condition)
    if hm:
        doc.add_paragraph()
        doc.add_picture(hm, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for var, name in [("TotalEV_MAX","Total EV"),("MaxBenthos","Benthos"),
                      ("ZooScore","Zooplankton"),("PhytoScore","Phytoplankton")]:
        if var in eva.columns:
            logger.info("  Generating %s map...", name)
            doc.add_heading(name, level=3)
            doc.add_picture(make_condition_map_matplotlib(overlay, eva, var, name), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 4. Supply
    doc.add_heading("4. Supply Table", level=1)
    sup = supply.rename(columns={"EUNIS_code":"EUNIS Code","EUNIS_name":"Habitat",
                                 "Fisheries_proxy":"Fisheries","FoodWeb_proxy":"Food Web",
                                 "PrimaryProd_proxy":"Primary Prod."})
    add_table(doc, sup)
    doc.add_page_break()

    # 5. Accounts
    doc.add_heading("5. BBT8 Accounts Summary", level=1)
    acct = accounts.copy()
    acct["area_m2"] = acct["area_m2"].apply(lambda x: f"{x:,.0f}")
    acct["Habitat_EV"] = acct["Habitat_EV"].round(2)
    acct["Habitat_confidence"] = acct["Habitat_confidence"].round(2)
    acct.columns = ["EUNIS Code","Habitat","Area (m2)","Habitat EV","Confidence"]
    add_table(doc, acct)

    # 6. Data Quality
    doc.add_heading("6. Data Quality", level=1)
    doc.add_paragraph(f"{len(missing)} subzones with data quality issues.")
    if not missing.empty:
        add_table(doc, missing.head(20))

    # References
    doc.add_heading("References", level=1)
    for ref in [
        "Franco A. & Amorim E. (2025) EVA Guidance. MARBEFES WP4.1.",
        "EMODnet (2023) EUSeaMap 2023.",
        "UN (2021) SEEA EA.",
        "Razinkovas-Baziukas A. et al. (2025) Lithuanian BBT5 EVA report.",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    # Footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = f"MARBEFES Physical Accounts (BBT8) — Lithuanian BBT5 | {now} | EVA v3.4.0"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(8); r.font.color.rgb = RGBColor(108,117,125)

    path = os.path.join(TUTORIAL_DIR, "PhysicalAccounts_BBT8_LithuanianBBT5.docx")
    doc.save(path)
    logger.info("Word report: %s", path)
    return path


def generate_bbt8_excel(overlay, eva, extent, condition, supply, accounts, missing):
    """Generate BBT8-format Excel workbook."""
    logger.info("Generating BBT8 Excel...")
    # main_values
    eva_sub = eva[["Subzone_ID","TotalEV_MAX","Confidence"]].copy() if "Subzone_ID" in eva.columns else pd.DataFrame()
    if "geometry" in eva_sub.columns:
        eva_sub = eva_sub.drop(columns="geometry")
    if not eva_sub.empty:
        mv = overlay[["Subzone_ID","dominant_EUNIS"]].merge(eva_sub, on="Subzone_ID", how="left")
        mv.columns = ["Subzone_ID","EUNIS_code","Habitat_EV","Habitat_confidence"]
    else:
        mv = overlay[["Subzone_ID","dominant_EUNIS"]].copy()
        mv.columns = ["Subzone_ID","EUNIS_code"]
        mv["Habitat_EV"] = np.nan; mv["Habitat_confidence"] = np.nan

    metadata = {
        "Report": "SEEA EA Physical Accounts (BBT8 format)",
        "BBT": "BBT5 — Lithuanian Baltic Sea coast and Curonian Lagoon",
        "EAA": "Lithuanian EEZ, Territorial Sea, Curonian Lagoon",
        "Year": "2017-2023",
        "Framework": "SEEA EA / MARBEFES WP4",
        "Generated": datetime.now().strftime("%Y-%m-%d"),
        "EUNIS Source": "EMODnet EUSeaMap 2023",
        "EVA Version": "MARBEFES EVA v3.4.0",
        "Contact": "Klaipeda University / MARBEFES",
    }

    buf = generate_bbt8_workbook(
        accounts=accounts, main_values=mv, extent=extent,
        condition=condition, supply=supply, metadata=metadata,
        missing_values=missing,
    )
    path = os.path.join(TUTORIAL_DIR, "PhysicalAccounts_BBT8_LithuanianBBT5.xlsx")
    with open(path, "wb") as f:
        f.write(buf.read())
    logger.info("BBT8 Excel: %s", path)
    return path


def main():
    os.makedirs(TUTORIAL_DIR, exist_ok=True)
    overlay, eva = load_all()
    extent, condition, supply, accounts, missing = compute_all(overlay, eva)

    html_path = generate_html_report(overlay, eva, extent, condition, supply, accounts, missing)
    docx_path = generate_docx_report(overlay, eva, extent, condition, supply, accounts, missing)
    xlsx_path = generate_bbt8_excel(overlay, eva, extent, condition, supply, accounts, missing)

    logger.info("\nDone. Three outputs:")
    logger.info("  HTML: %s", html_path)
    logger.info("  Word: %s", docx_path)
    logger.info("  Excel: %s", xlsx_path)


if __name__ == "__main__":
    main()
