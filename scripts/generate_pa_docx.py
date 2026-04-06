"""Generate Physical Accounts report as MS Word (.docx) for Lithuanian BBT5.

Produces a professional Word document with tables, charts (as images),
and maps (as static images).
"""
import logging
import os
import sys
import io
from datetime import datetime

import geopandas as gpd
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
import plotly.io as pio
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

EVA_FINAL = os.path.normpath(
    r"C:\Users\DELL\OneDrive - ku.lt\HORIZON_EUROPE\MARBEFES\EVA_FINAL"
)
CORRECTED = os.path.normpath(
    r"C:\Users\DELL\OneDrive - ku.lt\HORIZON_EUROPE\MARBEFES\EVA_FINAL_corrected"
)
OUTPUT_DIR = os.path.normpath(
    os.path.join(os.path.dirname(__file__), "..", "tutorial")
)

HABITAT_COLORS = {
    "Circalittoral sand": "#f4a460",
    "Circalittoral mud": "#8b7355",
    "Circalittoral coarse sediment": "#d2b48c",
    "Infralittoral rock and biogenic reef": "#2e8b57",
    "Infralittoral sand": "#ffe4b5",
    "Circalittoral rock and biogenic reef": "#006400",
}

INDICATOR_NAMES = {
    "AQ7_HABITATS": "Habitat Diversity",
    "ZooScore": "Zooplankton",
    "PhytoScore": "Phytoplankton",
    "MaxBenthos": "Benthos",
    "EVA_all_fish": "Fish",
}


def classify_eva(val):
    if pd.isna(val): return "No Data"
    if val <= 1: return "Very Low"
    if val <= 2: return "Low"
    if val <= 3: return "Medium"
    if val <= 4: return "High"
    return "Very High"


def load_data():
    logger.info("Loading data...")
    hab_src = os.path.join(EVA_FINAL, "habitats_EVA", "habitats_final.shp")
    habitats = gpd.read_file(hab_src)
    habitats = habitats[habitats["MSFD_broad"].notna()].copy()
    habitats["area_ha"] = habitats["Shape_Area"] / 10000.0

    combined_src = os.path.join(CORRECTED, "ALL4EVA_2025_fixed_geometries.gpkg")
    if not os.path.exists(combined_src):
        combined_src = os.path.join(EVA_FINAL, "ALL4EVA_2025_fixed_geometries.gpkg")
    eva = gpd.read_file(combined_src)

    return habitats, eva


def compute_all(habitats, eva):
    """Compute extent, condition, and supply tables."""
    # Extent
    extent = habitats.groupby("MSFD_broad").agg(
        polygon_count=("Shape_Area", "count"),
        total_area_ha=("area_ha", "sum"),
    ).reset_index().rename(columns={"MSFD_broad": "Habitat Type"})
    total = extent["total_area_ha"].sum()
    extent["pct_of_total"] = (extent["total_area_ha"] / total * 100).round(1)
    extent["total_area_ha"] = extent["total_area_ha"].round(1)
    extent["total_area_km2"] = (extent["total_area_ha"] / 100).round(1)
    extent = extent.sort_values("total_area_ha", ascending=False).reset_index(drop=True)

    # Condition via spatial join
    hab_proj = habitats.to_crs(eva.crs) if habitats.crs != eva.crs else habitats
    hab_c = hab_proj[["MSFD_broad", "geometry"]].copy()
    hab_c["geometry"] = hab_c.geometry.centroid
    score_cols = [c for c in INDICATOR_NAMES if c in eva.columns]
    joined = gpd.sjoin(hab_c, eva[score_cols + ["geometry"]], how="left", predicate="within")

    condition = joined.groupby("MSFD_broad")[score_cols].mean().round(3)
    condition = condition.reset_index().rename(columns={"MSFD_broad": "Habitat Type"})

    # Supply
    supply_cols = [c for c in ["EVA_all_fish", "ZooScore", "PhytoScore"] if c in joined.columns]
    supply = joined.groupby("MSFD_broad")[supply_cols].mean().round(3)
    supply = supply.reset_index().rename(columns={"MSFD_broad": "Habitat Type"})

    return extent, condition, supply, habitats


# --- Chart generation (as PNG bytes) ---

def make_extent_chart_png(extent):
    fig, ax = plt.subplots(figsize=(10, 5))
    colors = [HABITAT_COLORS.get(h, "#999") for h in extent["Habitat Type"]]
    bars = ax.barh(extent["Habitat Type"], extent["total_area_km2"], color=colors, edgecolor="#333", linewidth=0.5)
    ax.set_xlabel("Area (km2)", fontsize=12)
    ax.set_title("Ecosystem Extent by Habitat Type", fontsize=14, fontweight="bold", color="#006994")
    for bar, val in zip(bars, extent["total_area_km2"]):
        ax.text(bar.get_width() + 5, bar.get_y() + bar.get_height()/2,
                f"{val} km2", va="center", fontsize=10)
    ax.invert_yaxis()
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def make_condition_chart_png(condition):
    avail = {k: v for k, v in INDICATOR_NAMES.items() if k in condition.columns}
    x = np.arange(len(condition))
    width = 0.15
    fig, ax = plt.subplots(figsize=(12, 6))
    colors_list = ["#4caf50", "#2196f3", "#ff9800", "#9c27b0", "#f44336"]
    for i, (col, name) in enumerate(avail.items()):
        vals = condition[col].fillna(0)
        ax.bar(x + i * width, vals, width, label=name, color=colors_list[i % len(colors_list)])
    ax.set_xticks(x + width * len(avail) / 2)
    ax.set_xticklabels(condition["Habitat Type"], rotation=25, ha="right", fontsize=9)
    ax.set_ylabel("Condition Score (0-5)", fontsize=12)
    ax.set_ylim(0, 5.5)
    ax.set_title("Ecosystem Condition by Habitat Type", fontsize=14, fontweight="bold", color="#006994")
    ax.legend(loc="upper right", fontsize=9)
    ax.axhline(y=1, color="#ddd", linestyle="--", linewidth=0.5)
    ax.axhline(y=2, color="#ddd", linestyle="--", linewidth=0.5)
    ax.axhline(y=3, color="#ddd", linestyle="--", linewidth=0.5)
    ax.axhline(y=4, color="#ddd", linestyle="--", linewidth=0.5)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def make_condition_heatmap_png(condition):
    avail = {k: v for k, v in INDICATOR_NAMES.items() if k in condition.columns}
    data = condition[list(avail.keys())].values
    fig, ax = plt.subplots(figsize=(10, 5))
    im = ax.imshow(data, cmap="RdYlGn", vmin=0, vmax=5, aspect="auto")
    ax.set_xticks(range(len(avail)))
    ax.set_xticklabels(list(avail.values()), fontsize=10)
    ax.set_yticks(range(len(condition)))
    ax.set_yticklabels(condition["Habitat Type"], fontsize=9)
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            val = data[i, j]
            color = "white" if (val < 1.5 or val > 3.5) else "black"
            if not np.isnan(val):
                ax.text(j, i, f"{val:.2f}", ha="center", va="center", fontsize=10, color=color)
    plt.colorbar(im, ax=ax, label="Score (0-5)", shrink=0.8)
    ax.set_title("Condition Heatmap", fontsize=14, fontweight="bold", color="#006994")
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def make_habitat_map_png(habitats):
    gdf = habitats.to_crs(epsg=4326)
    fig, ax = plt.subplots(figsize=(10, 8))
    for hab_type, color in HABITAT_COLORS.items():
        subset = gdf[gdf["MSFD_broad"] == hab_type]
        if not subset.empty:
            subset.plot(ax=ax, color=color, edgecolor="#333", linewidth=0.3, label=hab_type)
    ax.set_title("Habitat Distribution — Lithuanian BBT5", fontsize=14, fontweight="bold", color="#006994")
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    patches = [mpatches.Patch(color=c, label=n) for n, c in HABITAT_COLORS.items()]
    ax.legend(handles=patches, loc="lower left", fontsize=8, framealpha=0.9)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def make_condition_map_png(habitats, eva, variable, title):
    hab_proj = habitats.to_crs(eva.crs) if habitats.crs != eva.crs else habitats
    hab_c = hab_proj[["MSFD_broad", "geometry"]].copy()
    hab_c["geometry"] = hab_c.geometry.centroid
    if variable not in eva.columns:
        return None
    joined = gpd.sjoin(hab_c, eva[[variable, "geometry"]], how="left", predicate="within")
    gdf = habitats.copy()
    gdf[variable] = joined[variable].values[:len(gdf)]
    gdf = gdf.to_crs(epsg=4326)

    fig, ax = plt.subplots(figsize=(10, 8))
    gdf.plot(column=variable, ax=ax, cmap="RdYlGn", vmin=0, vmax=5,
             edgecolor="#333", linewidth=0.3, legend=True,
             legend_kwds={"label": f"{title} (0-5)", "shrink": 0.7})
    ax.set_title(f"{title} — Lithuanian BBT5", fontsize=14, fontweight="bold", color="#006994")
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# --- Word document generation ---

def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, df, col_widths=None):
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "006994")

    # Data rows
    for i, (_, row) in enumerate(df.iterrows()):
        for j, col in enumerate(df.columns):
            cell = table.rows[i + 1].cells[j]
            val = row[col]
            cell.text = str(val) if pd.notna(val) else ""
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "F0F7FA")

    return table


def generate_docx(extent, condition, supply, habitats, eva):
    logger.info("Generating Word document...")
    doc = Document()
    now = datetime.now().strftime("%Y-%m-%d")

    # --- Styles ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # --- Title Page ---
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("SEEA EA Physical Accounts", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 105, 148)

    subtitle = doc.add_paragraph("Lithuanian BBT5 — Curonian Lagoon and Baltic Sea Coast")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 184, 212)

    meta = doc.add_paragraph(f"MARBEFES WP4 | {now}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(108, 117, 125)

    doc.add_paragraph()
    doc.add_paragraph("EU Horizon Europe Research Programme").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Table of Contents placeholder ---
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("2. Input Data and Methodology")
    doc.add_paragraph("3. Extent Account")
    doc.add_paragraph("4. Condition Account")
    doc.add_paragraph("5. Ecosystem Service Supply Table")
    doc.add_paragraph("6. Summary and Conclusions")
    doc.add_paragraph("References")
    doc.add_page_break()

    # --- 1. Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This report presents the Physical Natural Capital Accounts for the Lithuanian "
        "Broad Biotope Type 5 (BBT5), covering the Curonian Lagoon and Baltic Sea coast. "
        "The accounts follow the UN System of Environmental-Economic Accounting — Ecosystem "
        "Accounting (SEEA EA) framework as adapted for marine environments under the MARBEFES project."
    )
    doc.add_paragraph("Three account components are presented:")
    doc.add_paragraph("Extent Account — area (hectares) per benthic habitat type", style="List Bullet")
    doc.add_paragraph("Condition Account — ecological condition per habitat type using EVA scores", style="List Bullet")
    doc.add_paragraph("Supply Table — ecosystem service provision proxies where data permits", style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Ecosystem Accounting Area (EAA): ").bold = True
    p.add_run("Lithuanian EEZ, Territorial Sea, and Curonian Lagoon")
    p = doc.add_paragraph()
    p.add_run("Accounting period: ").bold = True
    p.add_run("2017-2023 (monitoring data period)")
    p = doc.add_paragraph()
    p.add_run("Spatial resolution: ").bold = True
    p.add_run("310 habitat polygons (HELCOM HUB Level 3 / EUNIS Level 2)")

    # --- 2. Input Data ---
    doc.add_heading("2. Input Data and Methodology", level=1)

    doc.add_heading("2.1 Habitat Classification", level=2)
    doc.add_paragraph(
        "Benthic habitats were classified according to HELCOM HUB Level 3 / EUNIS Level 2, "
        "derived from the Lithuanian Environmental Protection Agency's 2019 habitat mapping. "
        "Six broad habitat types were identified in the Baltic Sea portion of the EAA."
    )

    doc.add_heading("2.2 Data Sources", level=2)
    sources_df = pd.DataFrame({
        "Component": ["Habitats", "Benthos", "Fish", "Zooplankton", "Phytoplankton"],
        "Source": [
            "Lithuanian EPA (HELCOM HUB L3, 2019)",
            "SDM + video surveys (2021-2023)",
            "ICES BITS + EPA catch stats (2000-2023)",
            "MRI surveys + BIO-C3/RETRO (1996-2020)",
            "EPA monitoring (2017-2023)",
        ],
        "Coverage": ["Full EAA", "Baltic coast", "BS + CL", "21 stations", "18 stations"],
    })
    add_styled_table(doc, sources_df)

    doc.add_heading("2.3 Methodology", level=2)
    doc.add_paragraph(
        "Extent: Computed from polygon areas grouped by MSFD broad habitat type. "
        "Condition: EVA scores spatially joined to habitat polygons via centroid-in-hexagon matching. "
        "Supply: Fish CPUE and plankton scores used as ecosystem service proxies."
    )

    doc.add_page_break()

    # --- 3. Extent Account ---
    doc.add_heading("3. Extent Account", level=1)

    total_ha = extent["total_area_ha"].sum()
    total_km2 = extent["total_area_km2"].sum()
    doc.add_paragraph(
        f"The total marine habitat extent within the Lithuanian EAA is {total_ha:,.0f} hectares "
        f"({total_km2:,.0f} km2). Sandy substrates dominate, covering nearly 43% of the seabed."
    )

    doc.add_heading("3.1 Extent Table", level=2)
    extent_display = extent[["Habitat Type", "total_area_ha", "total_area_km2", "pct_of_total"]].copy()
    extent_display.columns = ["Habitat Type", "Area (Ha)", "Area (km2)", "% of Total"]
    # Add total row
    total_row = pd.DataFrame([{
        "Habitat Type": "TOTAL",
        "Area (Ha)": total_ha.round(1),
        "Area (km2)": total_km2.round(1),
        "% of Total": 100.0,
    }])
    extent_display = pd.concat([extent_display, total_row], ignore_index=True)
    add_styled_table(doc, extent_display)

    doc.add_heading("3.2 Extent Chart", level=2)
    logger.info("  Generating extent chart...")
    chart_buf = make_extent_chart_png(extent)
    doc.add_picture(chart_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.3 Habitat Distribution Map", level=2)
    logger.info("  Generating habitat map...")
    map_buf = make_habitat_map_png(habitats)
    doc.add_picture(map_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 1. Spatial distribution of benthic habitat types in the Lithuanian BBT5.").italic = True

    doc.add_page_break()

    # --- 4. Condition Account ---
    doc.add_heading("4. Condition Account", level=1)
    doc.add_paragraph(
        "Ecological condition is assessed using EVA scores (0-5 scale) across five indicators. "
        "Each habitat polygon is assigned the EVA score of the hexagonal grid cell containing its centroid."
    )

    p = doc.add_paragraph()
    p.add_run("Note on confidence: ").bold = True
    p.add_run(
        "All condition scores have Low confidence (each EC answered only 1-4 of 7-8 possible "
        "Assessment Questions). Scores should be interpreted as preliminary indicators."
    )

    doc.add_heading("4.1 Condition Summary", level=2)
    cond_display = condition.copy()
    for col in list(INDICATOR_NAMES.keys()):
        if col in cond_display.columns:
            nice = INDICATOR_NAMES[col]
            cond_display = cond_display.rename(columns={col: nice})
            cond_display[f"{nice} Class"] = cond_display[nice].apply(classify_eva)
    add_styled_table(doc, cond_display)

    doc.add_heading("4.2 Condition Comparison Chart", level=2)
    logger.info("  Generating condition chart...")
    chart_buf = make_condition_chart_png(condition)
    doc.add_picture(chart_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.3 Condition Heatmap", level=2)
    logger.info("  Generating condition heatmap...")
    hm_buf = make_condition_heatmap_png(condition)
    doc.add_picture(hm_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 2. Condition heatmap: habitat type vs. ecological indicator.").italic = True

    doc.add_heading("4.4 Condition Maps", level=2)
    for var, name in [("AQ7_HABITATS", "Habitat Diversity"),
                      ("MaxBenthos", "Benthos Condition"),
                      ("ZooScore", "Zooplankton Condition"),
                      ("PhytoScore", "Phytoplankton Condition")]:
        logger.info("  Generating %s map...", name)
        map_buf = make_condition_map_png(habitats, eva, var, name)
        if map_buf:
            doc.add_paragraph(f"{name}", style="Heading 3")
            doc.add_picture(map_buf, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- 5. Supply Table ---
    doc.add_heading("5. Ecosystem Service Supply Table", level=1)
    doc.add_paragraph(
        "The supply table presents available proxy indicators for ecosystem services per habitat type. "
        "Full SEEA EA compliance requires physical units (tonnes, m3, etc.); the values below are "
        "relative EVA scores (0-5) serving as provisional indicators."
    )

    doc.add_heading("5.1 Available Service Proxies", level=2)
    supply_display = supply.rename(columns={
        "EVA_all_fish": "Fisheries (proxy)",
        "ZooScore": "Food Web (proxy)",
        "PhytoScore": "Primary Prod. (proxy)",
    })
    add_styled_table(doc, supply_display)

    doc.add_heading("5.2 Data Gaps", level=2)
    doc.add_paragraph(
        "The following ecosystem services require additional data for full physical accounting:"
    )
    gaps_df = pd.DataFrame({
        "Service": ["Carbon Sequestration", "Coastal Protection", "Tourism & Recreation", "Nutrient Cycling"],
        "Status": ["NOT AVAILABLE"] * 4,
        "Units Needed": ["t C/ha/yr", "wave attenuation index", "visits/yr", "kg N-P/ha/yr"],
        "Potential Source": [
            "Sediment core studies, IPCC factors",
            "SHYFEM hydrodynamic model",
            "Tourism statistics, visitor surveys",
            "HELCOM PLC, biogeochemical models",
        ],
    })
    add_styled_table(doc, gaps_df)

    doc.add_page_break()

    # --- 6. Summary ---
    doc.add_heading("6. Summary and Conclusions", level=1)

    doc.add_heading("Key Findings", level=2)
    doc.add_paragraph(
        "Sandy substrates (circalittoral + infralittoral sand) dominate at 43% of total area. "
        "Rock and biogenic reef habitats cover only 9.6% but are the most ecologically diverse.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Zooplankton scores are Very High across all habitats (3.5-4.3), dominating the overall "
        "condition assessment. Habitat diversity is highest in infralittoral rock/reef (Medium, 2.3).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Fish provisioning data is available for only 3 of 6 habitat types. Carbon, tourism, "
        "and nutrient cycling data are entirely absent.",
        style="List Bullet"
    )

    doc.add_heading("Limitations", level=2)
    doc.add_paragraph("All condition scores have Low confidence due to limited AQ coverage", style="List Bullet")
    doc.add_paragraph("No temporal baseline — single-period snapshot, not trend analysis", style="List Bullet")
    doc.add_paragraph("Curonian Lagoon habitats not included (different classification)", style="List Bullet")
    doc.add_paragraph("Supply table uses proxy scores, not physical units per SEEA EA", style="List Bullet")

    doc.add_heading("Next Steps", level=2)
    doc.add_paragraph("Acquire fisheries catch data in physical units (tonnes/year)", style="List Number")
    doc.add_paragraph("Integrate satellite-derived primary production (Copernicus Marine)", style="List Number")
    doc.add_paragraph("Establish temporal baselines for condition trend monitoring", style="List Number")
    doc.add_paragraph("Extend to Curonian Lagoon using EUNIS Level 2 classification", style="List Number")
    doc.add_paragraph("Connect with ARIES platform for automated accounting", style="List Number")

    # --- References ---
    doc.add_heading("References", level=1)
    refs = [
        "Franco A. & Amorim E. (2025) Ecological Value Assessment (EVA) — Guidance including FAQs. MARBEFES WP4.1.",
        "Razinkovas-Baziukas A. et al. (2025) Curonian Lagoon and Baltic Sea coast Lithuanian BBT EVA report. Klaipeda University / MARBEFES.",
        "UN (2021) System of Environmental-Economic Accounting — Ecosystem Accounting (SEEA EA). United Nations.",
        "Luisetti T. & Burdon D. (2023) Draft Guidance on Socio-Economic Frameworks and Methods — Physical Accounts Section. MARBEFES D4.2.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    # --- Footer on all pages ---
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = f"MARBEFES Physical Accounts — Lithuanian BBT5 | Generated {now} | EVA v3.3.1"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(108, 117, 125)

    # Save
    output_path = os.path.join(OUTPUT_DIR, "PhysicalAccounts_Report_LithuanianBBT5.docx")
    doc.save(output_path)
    logger.info("Word document saved to %s", output_path)
    return output_path


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    habitats, eva = load_data()
    extent, condition, supply, habitats = compute_all(habitats, eva)
    path = generate_docx(extent, condition, supply, habitats, eva)
    logger.info("Done. Open: %s", path)


if __name__ == "__main__":
    main()
