# tests/test_pa_docx.py
"""Tests for pa_docx module (stateless DOCX report generator)."""
from __future__ import annotations

import io

import geopandas as gpd
import numpy as np
import pandas as pd
import pytest
from docx import Document
from docx.oxml.ns import qn
from shapely.geometry import box

import pa_docx


# ---------------------------------------------------------------------------
# Synthetic fixtures
# ---------------------------------------------------------------------------
def _extent_df():
    return pd.DataFrame({
        "EUNIS_code":   ["A5.25", "A4.4", "A5.23"],
        "EUNIS_name":   ["Circalittoral fine sand",
                          "Baltic exposed circalittoral rock",
                          "Infralittoral fine sand"],
        "n_subzones":   [2, 1, 1],
        "total_area":   [200.0, 50.0, 100.0],  # Ha
        "pct_of_total": [57.1, 14.3, 28.6],
        "area_m2":      [2_000_000, 500_000, 1_000_000],
    })


def _condition_df():
    return pd.DataFrame({
        "EUNIS_code":         ["A5.25", "A4.4", "A5.23"],
        "EUNIS_name":         ["Circalittoral fine sand",
                                "Baltic exposed circalittoral rock",
                                "Infralittoral fine sand"],
        "n_subzones":         [2, 1, 1],
        "Habitat_EV":         [3.5, 4.2, 2.1],
        "Habitat_confidence": [np.nan, np.nan, np.nan],
        "AQ7_HABITATS_avg":   [2.0, 3.0, 1.5],
        "ZooScore_avg":       [4.0, 3.5, 3.0],
        "PhytoScore_avg":     [2.0, 1.5, 3.0],
        "MaxBenthos_avg":     [2.5, 4.0, 1.5],
    })


def _supply_df():
    return pd.DataFrame({
        "EUNIS_code":         ["A5.25", "A4.4", "A5.23"],
        "EUNIS_name":         ["Circalittoral fine sand",
                                "Baltic exposed circalittoral rock",
                                "Infralittoral fine sand"],
        "Fisheries_proxy":    [1.0, 2.0, np.nan],
        "FoodWeb_proxy":      [4.0, 3.5, 3.0],
        "PrimaryProd_proxy":  [2.0, 1.5, 3.0],
    })


def _missing_df():
    return pd.DataFrame({
        "Subzone_ID": ["R003_C001", "R003_C002"],
        "issue_type": ["no_eunis", "low_coverage"],
        "notes":      ["No EUNIS attribution", "EUNIS coverage only 20%"],
    })


def _overlay_gdf():
    return gpd.GeoDataFrame({
        "Subzone_ID":           ["R001_C001", "R001_C002", "R002_C001", "R002_C002"],
        "dominant_EUNIS":       ["A5.25", "A5.25", "A4.4", "A5.23"],
        "dominant_EUNIS_name":  ["C fine sand", "C fine sand", "C rock", "I fine sand"],
        "coverage_pct":         [95.0, 100.0, 85.0, 90.0],
    }, geometry=[box(0, 0, 1, 1), box(1, 0, 2, 1), box(0, 1, 1, 2), box(1, 1, 2, 2)],
       crs="EPSG:3346")


def _eva_gdf():
    return gpd.GeoDataFrame({
        "Subzone_ID":    ["R001_C001", "R001_C002", "R002_C001", "R002_C002"],
        "TotalEV_MAX":   [4.0, 3.5, 4.5, 3.0],
        "AQ7_HABITATS":  [1.5, 2.0, 3.0, 1.0],
        "ZooScore":      [4.0, 3.5, 4.5, 3.0],
        "PhytoScore":    [2.0, 2.5, 1.5, 3.0],
        "MaxBenthos":    [3.0, 2.0, 4.0, 1.5],
    }, geometry=[box(0, 0, 1, 1), box(1, 0, 2, 1), box(0, 1, 1, 2), box(1, 1, 2, 2)],
       crs="EPSG:3346")


# ---------------------------------------------------------------------------
# classify_eva — boundary behaviour
# ---------------------------------------------------------------------------
class TestClassifyEva:
    def test_nan_is_no_data(self):
        assert pa_docx.classify_eva(np.nan) == "No Data"
        assert pa_docx.classify_eva(pd.NA) == "No Data"

    def test_boundaries_are_inclusive_on_upper(self):
        # classify_eva uses <= boundaries
        assert pa_docx.classify_eva(0) == "Very Low"
        assert pa_docx.classify_eva(1) == "Very Low"
        assert pa_docx.classify_eva(1.01) == "Low"
        assert pa_docx.classify_eva(2) == "Low"
        assert pa_docx.classify_eva(3) == "Medium"
        assert pa_docx.classify_eva(4) == "High"
        assert pa_docx.classify_eva(4.5) == "Very High"
        assert pa_docx.classify_eva(5) == "Very High"
        # Values above 5 still classify as Very High (defensive)
        assert pa_docx.classify_eva(7.5) == "Very High"


# ---------------------------------------------------------------------------
# Formatters
# ---------------------------------------------------------------------------
class TestFormatters:
    def test_fmt_int_rounds_and_commas(self):
        assert pa_docx._fmt_int(1234) == "1,234"
        assert pa_docx._fmt_int(1234.7) == "1,235"
        assert pa_docx._fmt_int(0) == "0"

    def test_fmt_int_nan_empty_string(self):
        assert pa_docx._fmt_int(np.nan) == ""

    def test_fmt_int_non_numeric_falls_back_to_str(self):
        assert pa_docx._fmt_int("abc") == "abc"

    def test_fmt_float_default_two_digits(self):
        assert pa_docx._fmt_float(3.14159) == "3.14"
        assert pa_docx._fmt_float(1234.5) == "1,234.50"

    def test_fmt_float_digits_override(self):
        assert pa_docx._fmt_float(3.14159, digits=4) == "3.1416"

    def test_fmt_float_nan_empty(self):
        assert pa_docx._fmt_float(np.nan) == ""


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------
class TestParseMarkdown:
    def test_headings(self):
        blocks = pa_docx.parse_markdown("# H1\n\n## H2\n\n### H3\n")
        kinds = [b.kind for b in blocks]
        texts = [b.text for b in blocks]
        assert kinds == ["h1", "h2", "h3"]
        assert texts == ["H1", "H2", "H3"]

    def test_paragraph_joins_wrapped_lines(self):
        md = "First line of a\nparagraph that wraps.\n"
        blocks = pa_docx.parse_markdown(md)
        assert len(blocks) == 1
        assert blocks[0].kind == "para"
        assert blocks[0].text == "First line of a paragraph that wraps."

    def test_single_line_bullet(self):
        blocks = pa_docx.parse_markdown("- just one\n")
        assert [b.kind for b in blocks] == ["bullet"]
        assert blocks[0].text == "just one"

    def test_multiline_bullet_consumes_continuation(self):
        md = (
            "- First bullet\n"
            "  continues here\n"
            "- Second bullet\n"
        )
        blocks = pa_docx.parse_markdown(md)
        assert [b.kind for b in blocks] == ["bullet", "bullet"]
        assert blocks[0].text == "First bullet continues here"
        assert blocks[1].text == "Second bullet"

    def test_bullet_terminates_on_blank_line(self):
        md = (
            "- Bullet one\n"
            "\n"
            "Paragraph after.\n"
        )
        blocks = pa_docx.parse_markdown(md)
        assert [b.kind for b in blocks] == ["bullet", "para"]

    def test_pipe_table(self):
        md = (
            "| Col1 | Col2 |\n"
            "|---|---:|\n"
            "| a | 1 |\n"
            "| b | 2 |\n"
        )
        blocks = pa_docx.parse_markdown(md)
        assert len(blocks) == 1
        tbl = blocks[0]
        assert tbl.kind == "table"
        assert tbl.rows[0] == ["Col1", "Col2"]
        assert tbl.rows[1] == ["a", "1"]
        assert tbl.rows[2] == ["b", "2"]
        assert tbl.aligns == ["l", "r"]

    def test_parse_table_align_center(self):
        assert pa_docx._parse_table_align("|:---:|:--|---:|") == ["c", "l", "r"]


# ---------------------------------------------------------------------------
# Inline-run rendering
# ---------------------------------------------------------------------------
class TestInlineRuns:
    def test_plain_single_run(self):
        doc = Document()
        p = doc.add_paragraph()
        pa_docx.add_inline_runs(p, "plain text")
        assert [r.text for r in p.runs] == ["plain text"]
        assert not p.runs[0].bold and not p.runs[0].italic

    def test_bold_and_italic_split(self):
        doc = Document()
        p = doc.add_paragraph()
        pa_docx.add_inline_runs(p, "start **bold** middle *ital* end")
        texts = [r.text for r in p.runs]
        # order preserved: plain, bold, plain, italic, plain
        assert texts == ["start ", "bold", " middle ", "ital", " end"]
        assert p.runs[1].bold is True
        assert p.runs[3].italic is True

    def test_inline_code_uses_monospace(self):
        doc = Document()
        p = doc.add_paragraph()
        pa_docx.add_inline_runs(p, "run `cmd arg` now")
        code_runs = [r for r in p.runs if r.text == "cmd arg"]
        assert code_runs, "code run missing"
        assert code_runs[0].font.name == "Consolas"


# ---------------------------------------------------------------------------
# Narrative builder
# ---------------------------------------------------------------------------
class TestNarrative:
    def test_contains_metadata_and_counts(self):
        md = pa_docx.build_narrative_md(
            overlay=_overlay_gdf(), eva=_eva_gdf(),
            extent=_extent_df(), condition=_condition_df(),
            missing=_missing_df(),
            metadata={"bbt_name": "Test BBT", "generated": "2026-04-22"},
        )
        assert "# Physical Accounts — Test BBT" in md
        assert "Generated 2026-04-22" in md
        # Headline table reflects actual counts
        assert "| EUNIS L3 classes identified | 3 |" in md
        assert "| Hexagonal subzones total | 4 |" in md

    def test_top_three_table_sorted_by_area(self):
        md = pa_docx.build_narrative_md(
            overlay=_overlay_gdf(), eva=_eva_gdf(),
            extent=_extent_df(), condition=_condition_df(),
            missing=_missing_df(),
            metadata={"bbt_name": "X", "generated": "d"},
        )
        # A5.25 has the largest total_area in the synthetic data
        idx_top = md.find("| 1 | A5.25 |")
        assert idx_top != -1, "A5.25 should be rank 1 by area"


# ---------------------------------------------------------------------------
# DOCX assembly (integration)
# ---------------------------------------------------------------------------
class TestBuildDocx:
    def test_returns_bytesio(self):
        md = "# Title\n\n## 2. Extent Account\n\nBody.\n"
        buf = pa_docx.build_docx_bytes(
            md=md,
            extent=_extent_df(), condition=_condition_df(),
            supply=_supply_df(), missing=_missing_df(),
            maps={},
            metadata={"bbt_name": "X", "generated": "d"},
        )
        assert isinstance(buf, io.BytesIO)
        assert buf.getvalue()[:4] == b"PK\x03\x04"  # zip magic

    def test_section_order_and_sectpr_terminal(self):
        """The key structural property: §2 before §3 before §4, sectPr last."""
        md = (
            "# Title\n\n"
            "*MARBEFES WP4 | Generated 2026-04-22*\n\n"
            "## 2. Extent Account\n\nExtent body.\n\n"
            "## 3. Condition Account\n\nCondition body.\n\n"
            "## 4. Supply Account (Proxy)\n\nSupply body.\n"
        )
        buf = pa_docx.build_docx_bytes(
            md=md,
            extent=_extent_df(), condition=_condition_df(),
            supply=_supply_df(), missing=_missing_df(),
            maps={},
            metadata={"bbt_name": "X", "generated": "2026-04-22"},
        )
        doc = Document(buf)

        # sectPr must be the last element in body (structural invariant)
        body = doc.element.body
        last = list(body.iterchildren())[-1]
        assert last.tag.endswith("sectPr"), \
            f"sectPr must be last, got {last.tag}"

        # Heading-1 order in the flow
        h1s = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        idx2 = h1s.index("2. Extent Account")
        idx3 = h1s.index("3. Condition Account")
        idx4 = h1s.index("4. Supply Account (Proxy)")
        assert idx2 < idx3 < idx4

    def test_detail_tables_appended_per_section(self):
        """Each recognised section gets its data-driven detail table injected."""
        md = (
            "## 2. Extent Account\n\nbody\n\n"
            "## 3. Condition Account\n\nbody\n\n"
            "## 4. Supply Account (Proxy)\n\nbody\n"
        )
        buf = pa_docx.build_docx_bytes(
            md=md,
            extent=_extent_df(), condition=_condition_df(),
            supply=_supply_df(), missing=_missing_df(),
            maps={},
            metadata={"bbt_name": "X", "generated": "d"},
        )
        doc = Document(buf)
        tables = doc.tables
        # Extent detail has 4 data rows (3 habitats + TOTAL) = 5 rows including header
        extent_tbls = [t for t in tables if t.rows[0].cells[0].text == "EUNIS"]
        assert len(extent_tbls) >= 3  # extent + condition + supply

        # Extent detail must contain a TOTAL row
        total_present = any(
            "TOTAL" in (t.rows[i].cells[0].text for i in range(len(t.rows)))
            for t in extent_tbls
        )
        assert total_present, "Extent detail must have a TOTAL row"

    def test_italic_preamble_subtitle_is_skipped(self):
        """The '*MARBEFES ... Generated ...*' line already appears on title page."""
        md = (
            "# Title\n\n"
            "*MARBEFES WP4 | Generated 2026-01-01*\n\n"
            "## 1. Overview\n\nOverview text.\n"
        )
        buf = pa_docx.build_docx_bytes(
            md=md,
            extent=pd.DataFrame(), condition=pd.DataFrame(),
            supply=pd.DataFrame(), missing=pd.DataFrame(),
            maps={},
            metadata={"bbt_name": "X", "generated": "2026-01-01"},
        )
        doc = Document(buf)
        # The italic subtitle from MD must not appear twice. Title page
        # uses "Generated 2026-01-01"; the MD-sourced italic is skipped.
        gen_hits = sum(
            1 for p in doc.paragraphs
            if "Generated 2026-01-01" in p.text
        )
        assert gen_hits == 1, f"Expected exactly one 'Generated …' line, got {gen_hits}"

    def test_empty_dataframes_graceful(self):
        """Runs on sparse inputs without raising."""
        md = "## 2. Extent Account\n\nbody\n"
        buf = pa_docx.build_docx_bytes(
            md=md,
            extent=pd.DataFrame(), condition=pd.DataFrame(),
            supply=pd.DataFrame(), missing=pd.DataFrame(),
            maps={},
            metadata={"bbt_name": "X", "generated": "d"},
        )
        doc = Document(buf)
        # Should include a "No extent data available." paragraph
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No extent data available" in all_text
