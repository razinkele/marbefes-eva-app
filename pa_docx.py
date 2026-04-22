"""MARBEFES Physical Accounts DOCX report generator.

Stateless, no Shiny deps. Accepts plain DataFrames + GeoDataFrames and
returns ``io.BytesIO`` holding a Word document.

Public API
----------
* :func:`render_bbt8_maps` — render the 7 standard PA maps as PNG BytesIO
* :func:`build_narrative_md` — produce the narrative Markdown as a string
* :func:`build_docx_bytes` — assemble a DOCX from MD + data tables + maps
* :func:`generate_bbt8_docx_report` — one-call orchestrator used by the app

The CLI ``scripts/render_pa_lt_docx.py`` is a thin wrapper around this
module.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Mapping

import geopandas as gpd
import matplotlib
matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------
BRAND_TEAL = RGBColor(0x00, 0x69, 0x94)
BRAND_CYAN = RGBColor(0x00, 0xB8, 0xD4)
MUTED = RGBColor(0x6C, 0x75, 0x7D)
HEADER_FILL = "006994"
BAND_FILL = "F0F7FA"

INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")

EUNIS_COLORS = {
    "A3.4":                     "#006400",
    "A4.4":                     "#2e8b57",
    "A5.13":                    "#deb887",
    "A5.14":                    "#d2b48c",
    "A5.23":                    "#ffe4b5",
    "A5.24 or A5.33 or A5.34":  "#cd853f",
    "A5.25":                    "#f4a460",
    "A5.26 or A5.35 or A5.36":  "#8b7355",
    "A5.27 or A5.37":           "#a0522d",
}

EV_CLASS_COLORS = {
    "No Data":   "#eeeeee",
    "Very Low":  "#d73027",
    "Low":       "#fc8d59",
    "Medium":    "#fee08b",
    "High":      "#91cf60",
    "Very High": "#1a9850",
}

MAP_TITLES = {
    "EUNIS_classes":      "Figure 1. EUNIS Level 3 habitat classes.",
    "habEV_classes":      "Figure 2. Per-habitat ecological value class (habEV), area-weighted from EVA TotalEV_MAX.",
    "TotalEV_MAX":        "Figure 3. Total Ecological Value (EVA TotalEV_MAX) per subzone.",
    "AQ7_Habitats":       "Figure 4. AQ7 — Habitat-forming species.",
    "Benthos_MAX":        "Figure 5. Benthos Assessment Question scores (max across AQ6/9/13).",
    "AQ_Zooplankton":     "Figure 6. Zooplankton Assessment Question score.",
    "AQ_Phytoplankton":   "Figure 7. Phytoplankton Assessment Question score.",
}

EXTENT_MAPS = ["EUNIS_classes"]
CONDITION_MAPS = [
    "habEV_classes",
    "TotalEV_MAX",
    "AQ7_Habitats",
    "Benthos_MAX",
    "AQ_Zooplankton",
    "AQ_Phytoplankton",
]


# ---------------------------------------------------------------------------
# Classification helper (shared with generate_pa_lt_report)
# ---------------------------------------------------------------------------
def classify_eva(val) -> str:
    if pd.isna(val):
        return "No Data"
    if val <= 1:
        return "Very Low"
    if val <= 2:
        return "Low"
    if val <= 3:
        return "Medium"
    if val <= 4:
        return "High"
    return "Very High"


# ---------------------------------------------------------------------------
# Map rendering — each returns io.BytesIO of a PNG at 150 DPI
# ---------------------------------------------------------------------------
def _fig_to_bytesio(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _map_eunis_classes(overlay: gpd.GeoDataFrame, title: str) -> io.BytesIO:
    gdf = overlay.to_crs(epsg=4326)
    fig, ax = plt.subplots(figsize=(10, 8))
    no_data = gdf[gdf["dominant_EUNIS"].isna()]
    if not no_data.empty:
        no_data.plot(ax=ax, color="#eeeeee", edgecolor="#bbb", linewidth=0.3, zorder=1)
    for code, color in EUNIS_COLORS.items():
        sub = gdf[gdf["dominant_EUNIS"] == code]
        if not sub.empty:
            sub.plot(ax=ax, color=color, edgecolor="#333", linewidth=0.3, zorder=2)
    # EUNIS codes outside the LT palette — colour with a fallback cmap.
    extras = gdf[gdf["dominant_EUNIS"].notna() & ~gdf["dominant_EUNIS"].isin(EUNIS_COLORS)]
    if not extras.empty:
        extras.plot(ax=ax, column="dominant_EUNIS", legend=False,
                    edgecolor="#333", linewidth=0.3, zorder=2, cmap="tab20")
    ax.set_title(title, fontsize=13, fontweight="bold", color="#006994")
    ax.set_xlabel("Longitude"); ax.set_ylabel("Latitude")
    patches = [mpatches.Patch(color=c, label=code) for code, c in EUNIS_COLORS.items()]
    patches.append(mpatches.Patch(color="#eeeeee", label="No data"))
    ax.legend(handles=patches, loc="lower left", fontsize=7, framealpha=0.9)
    plt.tight_layout()
    return _fig_to_bytesio(fig)


def _map_indicator(
    overlay: gpd.GeoDataFrame, eva: gpd.GeoDataFrame, column: str, title: str,
) -> io.BytesIO | None:
    """Per-hex choropleth of an EVA indicator (joined by Subzone_ID)."""
    if column not in eva.columns or "Subzone_ID" not in eva.columns:
        return None
    merged = overlay[["Subzone_ID", "geometry"]].merge(
        eva[["Subzone_ID", column]], on="Subzone_ID", how="left",
    )
    gdf = gpd.GeoDataFrame(merged, geometry="geometry", crs=overlay.crs).to_crs(epsg=4326)
    fig, ax = plt.subplots(figsize=(10, 8))
    gdf.plot(column=column, ax=ax, cmap="RdYlGn", vmin=0, vmax=5,
             edgecolor="#333", linewidth=0.2, legend=True,
             missing_kwds={"color": "#eeeeee"}, zorder=2,
             legend_kwds={"label": f"{title} (0-5)", "shrink": 0.7})
    ax.set_title(title, fontsize=13, fontweight="bold", color="#006994")
    ax.set_xlabel("Longitude"); ax.set_ylabel("Latitude")
    plt.tight_layout()
    return _fig_to_bytesio(fig)


def _map_habEV_classes(
    overlay: gpd.GeoDataFrame, condition: pd.DataFrame, title: str,
) -> io.BytesIO | None:
    """Dissolve overlay by dominant_EUNIS and colour by habEV class."""
    if "Habitat_EV" not in condition.columns:
        return None
    class_map = {
        row["EUNIS_code"]: classify_eva(row["Habitat_EV"])
        for _, row in condition.iterrows()
    }
    gdf = overlay.copy()
    gdf["habEV_class"] = gdf["dominant_EUNIS"].map(class_map).fillna("No Data")
    gdf = gdf.to_crs(epsg=4326)
    fig, ax = plt.subplots(figsize=(10, 8))
    for cls, color in EV_CLASS_COLORS.items():
        sub = gdf[gdf["habEV_class"] == cls]
        if not sub.empty:
            sub.plot(ax=ax, color=color, edgecolor="#333", linewidth=0.4, zorder=2)
    ax.set_title(title, fontsize=12, fontweight="bold", color="#006994")
    ax.set_xlabel("Longitude"); ax.set_ylabel("Latitude")
    patches = [mpatches.Patch(color=c, label=n) for n, c in EV_CLASS_COLORS.items()]
    ax.legend(handles=patches, loc="lower left", fontsize=9, framealpha=0.9)
    plt.tight_layout()
    return _fig_to_bytesio(fig)


def render_bbt8_maps(
    overlay: gpd.GeoDataFrame,
    eva: gpd.GeoDataFrame,
    condition: pd.DataFrame,
) -> dict[str, io.BytesIO]:
    """Produce all 7 standard PA maps as PNG BytesIO buffers.

    Only maps whose underlying column exists in ``eva`` are returned.
    """
    maps: dict[str, io.BytesIO] = {}
    eunis_buf = _map_eunis_classes(overlay, "EUNIS Level 3 Habitat Classes")
    if eunis_buf is not None:
        maps["EUNIS_classes"] = eunis_buf

    habev_buf = _map_habEV_classes(overlay, condition, "habEV Classes (area-weighted TotalEV per habitat)")
    if habev_buf is not None:
        maps["habEV_classes"] = habev_buf

    indicator_specs = [
        ("TotalEV_MAX",   "Total Ecological Value",        "TotalEV_MAX"),
        ("MaxBenthos",    "Benthos AQ (max)",              "Benthos_MAX"),
        ("ZooScore",      "Zooplankton AQ",                "AQ_Zooplankton"),
        ("PhytoScore",    "Phytoplankton AQ",              "AQ_Phytoplankton"),
        ("AQ7_HABITATS",  "AQ7 — Habitat-forming species", "AQ7_Habitats"),
    ]
    for column, title, key in indicator_specs:
        buf = _map_indicator(overlay, eva, column, title)
        if buf is not None:
            maps[key] = buf
    return maps


# ---------------------------------------------------------------------------
# Narrative Markdown builder (mirrors generate_pa_lt_report.write_narrative)
# ---------------------------------------------------------------------------
def build_narrative_md(
    overlay: gpd.GeoDataFrame,
    eva: gpd.GeoDataFrame,
    extent: pd.DataFrame,
    condition: pd.DataFrame,
    missing: pd.DataFrame,
    metadata: Mapping[str, str],
) -> str:
    """Produce the Markdown narrative for the PA report."""
    bbt_name = metadata.get("bbt_name") or metadata.get("eaa_name") or "Ecosystem Accounting Area"
    generated = metadata.get("generated") or datetime.now().strftime("%Y-%m-%d")
    total_area_ha_col = "total_area" if "total_area" in extent.columns else "area_Ha"
    total_area_ha = float(extent[total_area_ha_col].sum()) if not extent.empty else 0.0
    n_habitats = int(len(extent))
    with_ev = int(condition["Habitat_EV"].notna().sum()) if "Habitat_EV" in condition.columns else 0

    md = f"""# Physical Accounts — {bbt_name}

*MARBEFES WP4 | SEEA EA Framework | Generated {generated}*

## 1. Overview

This report presents ecosystem **extent**, **condition**, and **supply**
accounts for {bbt_name}, following the SEEA Ecosystem Accounting framework
and the MARBEFES WP4 guidance (Luisetti & Burdon, 2023). Habitat
classification is **EUNIS Level 3**, derived from EMODnet EUSeaMap 2023 and
pre-joined to the hexagonal subzone grid. Ecological Value scores come
from the MARBEFES EVA pipeline.

**Headline figures**

| Metric | Value |
|---|---:|
| EUNIS L3 classes identified | {n_habitats} |
| Hexagonal subzones total | {len(overlay)} |
| Subzones with EUNIS attribution | {int(overlay['dominant_EUNIS'].notna().sum())} |
| Habitats with habEV computed | {with_ev} / {n_habitats} |
| Total mapped extent | {total_area_ha:,.0f} Ha |
| EVA source features | {len(eva):,} |

"""

    # Top three habitats — accommodate either "total_area" or "area_Ha"
    top = extent.sort_values(total_area_ha_col, ascending=False).head(3)
    cond_ev = condition.set_index("EUNIS_code")["Habitat_EV"] if "Habitat_EV" in condition.columns else pd.Series(dtype=float)
    md += "**Top three habitats by area**\n\n"
    md += "| Rank | EUNIS | Name | Area (Ha) | % | habEV |\n"
    md += "|---:|---|---|---:|---:|---:|\n"
    for i, (_, row) in enumerate(top.iterrows(), start=1):
        code = row["EUNIS_code"]
        name = row.get("EUNIS_name", code)
        pct = row.get("pct_of_total", np.nan)
        habev_val = cond_ev.get(code, np.nan)
        habev_txt = f"{habev_val:.2f}" if pd.notna(habev_val) else "n/a"
        pct_txt = f"{pct:.1f}" if pd.notna(pct) else "—"
        md += (f"| {i} | {code} | {name} | "
               f"{row[total_area_ha_col]:,.0f} | {pct_txt} | {habev_txt} |\n")

    md += f"""
## 2. Extent Account

Per-habitat area from dominant-EUNIS assignment on the hexagonal subzone
grid. See the *extent* sheet of the accompanying workbook for the full
table and the `area_m2` column for unrounded totals.

## 3. Condition Account

Condition is expressed through per-habitat **habEV**, an area-weighted mean
of the MARBEFES EVA TotalEV_MAX aggregated score (0-5 scale). The
`habEV_class` categorical bin (Very Low → Very High) is visualised below.

## 4. Supply Account (Proxy)

Three ecosystem-service proxies derived from EVA scores, per EUNIS class:

- **Fisheries_proxy** — EVA_all_fish (0-5 scale)
- **FoodWeb_proxy** — ZooScore (0-5 scale)
- **PrimaryProd_proxy** — PhytoScore (0-5 scale)

Full SEEA EA supply accounting in physical units (tonnes of fish, tCO2eq,
visitor-days, tonnes N removed, Ha protected) is flagged as *future work*.

## 5. Data Quality

{len(missing)} subzone-level issues were recorded — see the *missing_values*
sheet of the accompanying workbook for the full list.

## Methodology

- **Framework:** SEEA Ecosystem Accounting (UN, 2021).
- **Guidance:** MARBEFES WP4.3 Deliverable D4.2 (Luisetti & Burdon, 2023).
- **Habitat classification:** EUNIS Level 3 (EMODnet EUSeaMap 2023).
- **Ecological Value:** MARBEFES EVA aggregated score (0-5), area-weighted
  to habitat level as *habEV*.

## References

- Luisetti T., Burdon D. et al. (2023). *Draft Guidance on Socio-Economic Frameworks and Methods — Physical Accounts Section.* MARBEFES D4.2.
- UN (2021). *System of Environmental-Economic Accounting — Ecosystem Accounting (SEEA EA).*
- EMODnet (2023). *EUSeaMap 2023.*
- Franco A. & Amorim E. (2025). *EVA guidance.* MARBEFES WP4.1.
"""
    return md


# ---------------------------------------------------------------------------
# Markdown mini-parser (ported verbatim from scripts/render_pa_lt_docx.py)
# ---------------------------------------------------------------------------
@dataclass
class MdBlock:
    kind: str  # h1/h2/h3/para/bullet/table
    level: int = 0
    text: str = ""
    rows: list[list[str]] | None = None
    aligns: list[str] | None = None


def _parse_table_align(sep_row: str) -> list[str]:
    aligns = []
    for cell in [c.strip() for c in sep_row.strip().strip("|").split("|")]:
        if cell.startswith(":") and cell.endswith(":"):
            aligns.append("c")
        elif cell.endswith(":"):
            aligns.append("r")
        else:
            aligns.append("l")
    return aligns


def parse_markdown(md: str) -> list[MdBlock]:
    blocks: list[MdBlock] = []
    lines = md.splitlines()
    i = 0
    para_buf: list[str] = []

    def flush_para():
        if para_buf:
            text = " ".join(l.strip() for l in para_buf).strip()
            if text:
                blocks.append(MdBlock(kind="para", text=text))
            para_buf.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_para()
            i += 1
            continue

        if stripped.startswith("# "):
            flush_para()
            blocks.append(MdBlock(kind="h1", level=1, text=stripped[2:].strip()))
            i += 1
            continue
        if stripped.startswith("## "):
            flush_para()
            blocks.append(MdBlock(kind="h2", level=2, text=stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith("### "):
            flush_para()
            blocks.append(MdBlock(kind="h3", level=3, text=stripped[4:].strip()))
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            flush_para()
            buf = [stripped[2:].strip()]
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    break
                if nxt.startswith(("- ", "* ", "# ", "## ", "### ")) or nxt.startswith("|"):
                    break
                buf.append(nxt)
                j += 1
            blocks.append(MdBlock(kind="bullet", text=" ".join(buf)))
            i = j
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", nxt):
                flush_para()
                header_cells = [c.strip() for c in stripped.strip("|").split("|")]
                aligns = _parse_table_align(nxt)
                rows = [header_cells]
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                    rows.append(cells)
                    j += 1
                blocks.append(MdBlock(kind="table", rows=rows, aligns=aligns))
                i = j
                continue

        para_buf.append(line)
        i += 1

    flush_para()
    return blocks


# ---------------------------------------------------------------------------
# DOCX styling primitives
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_inline_runs(paragraph, text: str) -> None:
    tokens: list[tuple[str, str]] = []
    remaining = text
    while remaining:
        m_b = BOLD_RE.search(remaining)
        m_i = ITALIC_RE.search(remaining)
        m_c = INLINE_CODE_RE.search(remaining)
        candidates = [m for m in (m_b, m_i, m_c) if m]
        if not candidates:
            tokens.append(("plain", remaining))
            break
        m = min(candidates, key=lambda x: x.start())
        if m.start() > 0:
            tokens.append(("plain", remaining[: m.start()]))
        if m is m_b:
            tokens.append(("bold", m.group(1)))
        elif m is m_i:
            tokens.append(("italic", m.group(1)))
        else:
            tokens.append(("code", m.group(1)))
        remaining = remaining[m.end():]

    for kind, val in tokens:
        run = paragraph.add_run(val)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def _is_numeric_like(s: str) -> bool:
    s = s.replace(",", "").replace("%", "").strip()
    if not s:
        return False
    try:
        float(s)
        return True
    except ValueError:
        return False


def add_styled_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    aligns: list[str] | None = None,
    total_row_idx: int | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    if aligns is None:
        aligns = []
        probe = rows[0] if rows else []
        for j in range(len(headers)):
            cell_val = probe[j] if j < len(probe) else ""
            aligns.append("r" if _is_numeric_like(cell_val) else "l")

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, HEADER_FILL)

    align_map = {
        "l": WD_ALIGN_PARAGRAPH.LEFT,
        "r": WD_ALIGN_PARAGRAPH.RIGHT,
        "c": WD_ALIGN_PARAGRAPH.CENTER,
    }
    for i, row in enumerate(rows):
        is_total = (total_row_idx is not None and i == total_row_idx)
        for j in range(len(headers)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = align_map.get(aligns[j] if j < len(aligns) else "l", WD_ALIGN_PARAGRAPH.LEFT)
            val = row[j] if j < len(row) else ""
            run = p.add_run(val)
            run.font.size = Pt(9)
            if is_total:
                run.bold = True
                set_cell_shading(cell, "D9ECF2")
            elif i % 2 == 0:
                set_cell_shading(cell, BAND_FILL)


def add_figure_from_bytes(doc, image_bytes: io.BytesIO, caption: str, width_in: float = 6.0) -> None:
    image_bytes.seek(0)
    doc.add_picture(image_bytes, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


# ---------------------------------------------------------------------------
# Data-driven detail tables
# ---------------------------------------------------------------------------
def _fmt_int(v) -> str:
    if pd.isna(v):
        return ""
    try:
        return f"{int(round(float(v))):,}"
    except (ValueError, TypeError):
        return str(v)


def _fmt_float(v, digits: int = 2) -> str:
    if pd.isna(v):
        return ""
    try:
        return f"{float(v):,.{digits}f}"
    except (ValueError, TypeError):
        return str(v)


def _area_col(df: pd.DataFrame) -> str:
    for c in ("total_area", "area_Ha", "area"):
        if c in df.columns:
            return c
    return df.select_dtypes(include="number").columns[0]


def add_extent_detail(doc, extent: pd.DataFrame) -> None:
    doc.add_heading("Detailed extent table", level=2)
    if extent.empty:
        doc.add_paragraph("No extent data available.")
        return
    area_col = _area_col(extent)
    headers = ["EUNIS", "Habitat", "Subzones", "Area (Ha)", "Area (km²)", "% of Total"]
    total_ha = float(extent[area_col].sum())
    rows = []
    for _, r in extent.iterrows():
        rows.append([
            str(r["EUNIS_code"]),
            str(r.get("EUNIS_name", "")),
            _fmt_int(r.get("n_subzones", np.nan)),
            _fmt_float(r[area_col], 1),
            _fmt_float(r[area_col] / 100.0, 1),
            _fmt_float(r.get("pct_of_total", np.nan), 1),
        ])
    rows.append([
        "TOTAL", "", _fmt_int(extent.get("n_subzones", pd.Series([0])).sum()),
        _fmt_float(total_ha, 1), _fmt_float(total_ha / 100.0, 1), "100.0",
    ])
    add_styled_table(doc, headers, rows, aligns=["l", "l", "r", "r", "r", "r"],
                     total_row_idx=len(rows) - 1)


def add_condition_detail(doc, condition: pd.DataFrame) -> None:
    doc.add_heading("Detailed condition table", level=2)
    if condition.empty:
        doc.add_paragraph("No condition data available.")
        return
    headers = ["EUNIS", "Habitat", "habEV", "Class", "AQ7 Habitats",
               "Benthos (max)", "Zooplankton", "Phytoplankton"]
    rows = []
    for _, r in condition.iterrows():
        ev = r.get("Habitat_EV", np.nan)
        rows.append([
            str(r["EUNIS_code"]),
            str(r.get("EUNIS_name", "")),
            _fmt_float(ev, 2),
            classify_eva(ev),
            _fmt_float(r.get("AQ7_HABITATS_avg", np.nan), 2),
            _fmt_float(r.get("MaxBenthos_avg", np.nan), 2),
            _fmt_float(r.get("ZooScore_avg", np.nan), 2),
            _fmt_float(r.get("PhytoScore_avg", np.nan), 2),
        ])
    add_styled_table(doc, headers, rows, aligns=["l", "l", "r", "c", "r", "r", "r", "r"])


def add_supply_detail(doc, supply: pd.DataFrame) -> None:
    doc.add_heading("Detailed supply proxies", level=2)
    if supply.empty:
        doc.add_paragraph("No supply data available.")
        return
    headers = ["EUNIS", "Habitat", "Fisheries", "Food Web", "Primary Prod."]
    rows = []
    for _, r in supply.iterrows():
        rows.append([
            str(r["EUNIS_code"]),
            str(r.get("EUNIS_name", "")),
            _fmt_float(r.get("Fisheries_proxy", np.nan), 2),
            _fmt_float(r.get("FoodWeb_proxy", np.nan), 2),
            _fmt_float(r.get("PrimaryProd_proxy", np.nan), 2),
        ])
    add_styled_table(doc, headers, rows, aligns=["l", "l", "r", "r", "r"])


def add_missing_summary(doc, missing: pd.DataFrame) -> None:
    if missing is None or missing.empty:
        doc.add_paragraph("No missing-value issues detected.")
        return
    counts = missing["issue_type"].value_counts()
    add_styled_table(doc, ["Issue", "Hexes affected"],
                     [[str(k), _fmt_int(v)] for k, v in counts.items()],
                     aligns=["l", "r"])


# ---------------------------------------------------------------------------
# Section-grouped renderer
# ---------------------------------------------------------------------------
def _render_body_block(doc, block: MdBlock) -> None:
    if block.kind == "para":
        if block.text.startswith("*") and block.text.endswith("*") and not block.text.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(block.text.strip("*"))
            run.italic = True
            run.font.color.rgb = MUTED
            return
        p = doc.add_paragraph()
        add_inline_runs(p, block.text)
        return
    if block.kind == "bullet":
        p = doc.add_paragraph(style="List Bullet")
        add_inline_runs(p, block.text)
        return
    if block.kind == "h3":
        doc.add_heading(block.text, level=2)
        return
    if block.kind == "table":
        rows = block.rows or []
        if not rows:
            return
        headers = [re.sub(r"`", "", h) for h in rows[0]]
        data = [[re.sub(r"`", "", c) for c in r] for r in rows[1:]]
        aligns = block.aligns or ["l"] * len(headers)
        add_styled_table(doc, headers, data, aligns=aligns)


def _add_section_extras(
    doc,
    section_name: str,
    extent: pd.DataFrame,
    condition: pd.DataFrame,
    supply: pd.DataFrame,
    missing: pd.DataFrame,
    maps: Mapping[str, io.BytesIO],
) -> None:
    key = section_name.lower().strip()
    if key.startswith("2. extent"):
        add_extent_detail(doc, extent)
        for name in EXTENT_MAPS:
            if name in maps:
                add_figure_from_bytes(doc, maps[name], MAP_TITLES.get(name, name))
    elif key.startswith("3. condition"):
        add_condition_detail(doc, condition)
        for name in CONDITION_MAPS:
            if name in maps:
                add_figure_from_bytes(doc, maps[name], MAP_TITLES.get(name, name))
    elif key.startswith("4. supply"):
        add_supply_detail(doc, supply)
    elif key.startswith("5. data quality"):
        add_missing_summary(doc, missing)


def _render_grouped(
    doc,
    blocks: list[MdBlock],
    extent: pd.DataFrame,
    condition: pd.DataFrame,
    supply: pd.DataFrame,
    missing: pd.DataFrame,
    maps: Mapping[str, io.BytesIO],
) -> None:
    preamble: list[MdBlock] = []
    sections: list[tuple[str, list[MdBlock]]] = []
    current_name: str | None = None
    current_blocks: list[MdBlock] = []

    for b in blocks:
        if b.kind == "h1":
            continue
        if b.kind == "h2":
            if current_name is None:
                preamble = current_blocks
            else:
                sections.append((current_name, current_blocks))
            current_name = b.text
            current_blocks = []
        else:
            current_blocks.append(b)
    if current_name is not None:
        sections.append((current_name, current_blocks))
    else:
        preamble = current_blocks

    for b in preamble:
        # Skip the italic "*... Generated ...*" subtitle; it's on the title page.
        if b.kind == "para" and b.text.startswith("*") and b.text.endswith("*") \
                and "Generated" in b.text:
            continue
        _render_body_block(doc, b)

    for name, sub_blocks in sections:
        heading = doc.add_heading(name, level=1)
        for run in heading.runs:
            run.font.color.rgb = BRAND_TEAL
        for b in sub_blocks:
            _render_body_block(doc, b)
        _add_section_extras(doc, name, extent, condition, supply, missing, maps)


# ---------------------------------------------------------------------------
# Title / footer
# ---------------------------------------------------------------------------
def _add_title_page(doc, bbt_name: str, generated: str) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("SEEA EA Physical Accounts", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BRAND_TEAL

    subtitle = doc.add_paragraph(bbt_name)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = BRAND_CYAN

    meta = doc.add_paragraph(f"MARBEFES WP4 | SEEA EA Framework | Generated {generated}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = MUTED

    doc.add_paragraph()
    hosted = doc.add_paragraph("EU Horizon Europe Research Programme")
    hosted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hosted.runs[0].italic = True
    hosted.runs[0].font.color.rgb = MUTED

    doc.add_page_break()


def _add_footer(doc, bbt_name: str, generated: str) -> None:
    section = doc.sections[0]
    p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"MARBEFES Physical Accounts — {bbt_name}   |   Generated {generated}")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


# ---------------------------------------------------------------------------
# Public assembly API
# ---------------------------------------------------------------------------
def build_docx_bytes(
    md: str,
    extent: pd.DataFrame,
    condition: pd.DataFrame,
    supply: pd.DataFrame,
    missing: pd.DataFrame,
    maps: Mapping[str, io.BytesIO],
    metadata: Mapping[str, str],
) -> io.BytesIO:
    """Assemble a DOCX from MD narrative + data tables + in-memory maps."""
    blocks = parse_markdown(md)

    bbt_name = metadata.get("bbt_name") or metadata.get("eaa_name") or "Ecosystem Accounting Area"
    generated = metadata.get("generated") or datetime.now().strftime("%Y-%m-%d")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    _add_title_page(doc, bbt_name, generated)
    _render_grouped(doc, blocks, extent, condition, supply, missing, maps)
    _add_footer(doc, bbt_name, generated)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def generate_bbt8_docx_report(
    overlay: gpd.GeoDataFrame,
    eva: gpd.GeoDataFrame,
    extent: pd.DataFrame,
    condition: pd.DataFrame,
    supply: pd.DataFrame,
    missing: pd.DataFrame,
    metadata: Mapping[str, str],
) -> io.BytesIO:
    """One-call orchestrator: maps + narrative + DOCX assembly.

    ``overlay`` and ``eva`` are the raw GeoDataFrames used for per-hex maps.
    ``extent`` / ``condition`` / ``supply`` / ``missing`` are the BBT8-format
    DataFrames produced by :mod:`eunis_data`.
    """
    logger.info("Rendering PA maps...")
    maps = render_bbt8_maps(overlay, eva, condition)
    logger.info("Building narrative...")
    md = build_narrative_md(overlay, eva, extent, condition, missing, metadata)
    logger.info("Assembling DOCX...")
    return build_docx_bytes(md, extent, condition, supply, missing, maps, metadata)
